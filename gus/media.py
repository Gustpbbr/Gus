import base64
import asyncio
import hashlib
import logging
import os
from io import BytesIO

import httpx
from PIL import Image

logger = logging.getLogger(__name__)

# Vision recomenda até ~1568x1568 (1.15MP). Acima disso o Claude faz resize mesmo.
MAX_IMAGE_PIXELS = 1568 * 1568
JPEG_QUALITY = 85

# Claude native PDF: até 100 páginas, 32MB por arquivo
PDF_MAX_SIZE_MB = 32

# Whisper: até 25MB por arquivo
AUDIO_MAX_SIZE_MB = 25

# Cache por hash SHA-256 do arquivo — evita re-processar a mesma mídia.
#
# Budget: count + bytes. PDF/imagem otimizada chega em base64 (~33% maior
# que original). Container Railway pequeno, OOM por cache cheio é risco
# real — defesa: limite total de bytes além do count.
_media_cache: dict[str, tuple[list[dict], int]] = {}  # hash → (content, size_bytes)
_media_cache_bytes = 0  # soma corrente de bytes
_audio_cache: dict[str, str] = {}  # transcrições cacheadas separadamente
_CACHE_MAX_ITEMS = 50  # defesa contra fragmentação (PDFs minúsculos)
_CACHE_MAX_BYTES = 200 * 1024 * 1024  # 200MB total


def _content_bytes(content: list[dict]) -> int:
    """Aproxima bytes ocupados por um content (Vision/PDF blocks).

    Para blocos com `source.data` em base64, calcula tamanho do binário
    decodificado (len(base64) * 3/4). Inclui também tamanho de texto
    associado.
    """
    total = 0
    for block in content:
        if not isinstance(block, dict):
            continue
        source = block.get("source") or {}
        data = source.get("data") or ""
        if data:
            total += (len(data) * 3) // 4
        text = block.get("text") or ""
        total += len(text.encode("utf-8")) if text else 0
    return total


def _cache_get(file_hash: str) -> list[dict] | None:
    entry = _media_cache.get(file_hash)
    return entry[0] if entry else None


def _cache_put(file_hash: str, content: list[dict]) -> None:
    """Insere no cache com ejeção LRU por count OU bytes.

    Ejeção: enquanto adicionar este item ultrapassar count_max ou
    bytes_max, remove o mais antigo (insertion order). Item único maior
    que o budget é aceito mesmo assim (cache fica só com ele) — PDFs até
    32MB cabem confortável em 200MB.
    """
    global _media_cache_bytes
    size = _content_bytes(content)

    while _media_cache and (
        len(_media_cache) >= _CACHE_MAX_ITEMS
        or _media_cache_bytes + size > _CACHE_MAX_BYTES
    ):
        oldest_key = next(iter(_media_cache))
        _, oldest_size = _media_cache.pop(oldest_key)
        _media_cache_bytes -= oldest_size

    _media_cache[file_hash] = (content, size)
    _media_cache_bytes += size


async def _download(url: str) -> bytes:
    """Download com retry exponencial (2s, 4s, 8s) em falhas transitórias."""
    tentativas = 3
    ultima_excecao: Exception | None = None
    for i in range(tentativas):
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, timeout=60)
                response.raise_for_status()
                return response.content
        except (httpx.TimeoutException, httpx.NetworkError, httpx.HTTPStatusError) as e:
            ultima_excecao = e
            if i < tentativas - 1:
                espera = 2 ** (i + 1)
                logger.warning(f"Download falhou (tentativa {i+1}/{tentativas}): {e}. Aguardando {espera}s.")
                await asyncio.sleep(espera)
    raise ultima_excecao if ultima_excecao else RuntimeError("download falhou")


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


def _otimizar_imagem(data: bytes) -> tuple[bytes, str]:
    """Decodifica imagem, converte pra RGB, faz resize se exceder MAX_IMAGE_PIXELS,
    re-encoda como JPEG quality 85. Retorna (bytes, media_type)."""
    img = Image.open(BytesIO(data))

    # Converte pra RGB (remove alpha, CMYK, paleta)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # Resize preservando aspect ratio
    pixels = img.width * img.height
    if pixels > MAX_IMAGE_PIXELS:
        ratio = (MAX_IMAGE_PIXELS / pixels) ** 0.5
        novo_w = int(img.width * ratio)
        novo_h = int(img.height * ratio)
        img = img.resize((novo_w, novo_h), Image.LANCZOS)
        logger.info(f"Imagem redimensionada de {img.width}x{img.height} original → {novo_w}x{novo_h}")

    buf = BytesIO()
    img.save(buf, format="JPEG", quality=JPEG_QUALITY, optimize=True)
    return buf.getvalue(), "image/jpeg"


async def processar_imagem(file_url: str, caption: str = "") -> list[dict]:
    """Baixa, otimiza e encoda imagem pra Claude Vision. Cache por hash."""
    data = await _download(file_url)
    file_hash = _hash_bytes(data)

    cached = _cache_get(file_hash)
    if cached:
        logger.info(f"Cache hit imagem {file_hash}")
        bloco_imagem = cached
    else:
        bytes_otimizados, media_type = await asyncio.to_thread(_otimizar_imagem, data)
        b64 = base64.standard_b64encode(bytes_otimizados).decode("utf-8")
        bloco_imagem = [{
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": b64}
        }]
        _cache_put(file_hash, bloco_imagem)

    content = list(bloco_imagem)
    if caption:
        content.append({"type": "text", "text": caption})
    else:
        content.append({"type": "text", "text": "O que está nessa imagem? Descreva e comente o que for relevante."})
    return content


async def processar_pdf(file_url: str, caption: str = "") -> list[dict]:
    """Usa o suporte nativo do Claude a PDF (content type 'document') — texto + OCR
    automático, preserva layout. Até 100 páginas, 32MB."""
    data = await _download(file_url)
    tamanho_mb = len(data) / 1024 / 1024

    if tamanho_mb > PDF_MAX_SIZE_MB:
        return [{
            "type": "text",
            "text": (
                f"PDF tem {tamanho_mb:.1f}MB — acima do limite de {PDF_MAX_SIZE_MB}MB do Claude. "
                "Envie uma versão menor ou divida o arquivo."
            )
        }]

    file_hash = _hash_bytes(data)
    cached = _cache_get(file_hash)
    if cached:
        logger.info(f"Cache hit PDF {file_hash}")
        bloco_doc = cached
    else:
        b64 = base64.standard_b64encode(data).decode("utf-8")
        bloco_doc = [{
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64}
        }]
        _cache_put(file_hash, bloco_doc)

    content = list(bloco_doc)
    prompt = caption or "Analise este documento."
    content.append({"type": "text", "text": prompt})
    return content


async def processar_docx(file_url: str, caption: str = "") -> list[dict]:
    """Extrai texto de .docx (Word) e envia pra Claude como texto."""
    from docx import Document  # import lazy

    data = await _download(file_url)
    doc = Document(BytesIO(data))

    linhas = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            linhas.append(t)
    # Também extrai texto de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                linhas.append(" | ".join(celulas))

    texto = "\n\n".join(linhas) if linhas else "(documento sem texto extraível)"
    prompt = caption or "Analise este documento Word."
    return [{"type": "text", "text": f"{prompt}\n\n---\n\n{texto}"}]


async def transcrever_audio(file_url: str) -> str:
    """Baixa áudio do Telegram e transcreve via Whisper API (pt-BR).
    Retorna o texto transcrito ou mensagem de erro prefixada por '(erro...'"""
    data = await _download(file_url)
    tamanho_mb = len(data) / 1024 / 1024

    if tamanho_mb > AUDIO_MAX_SIZE_MB:
        return f"(áudio de {tamanho_mb:.1f}MB excede limite de {AUDIO_MAX_SIZE_MB}MB do Whisper)"

    file_hash = _hash_bytes(data)
    if file_hash in _audio_cache:
        logger.info(f"Cache hit áudio {file_hash}")
        return _audio_cache[file_hash]

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "(OPENAI_API_KEY não configurada no Railway — Whisper indisponível)"

    files = {
        "file": ("audio.ogg", data, "audio/ogg"),
        "model": (None, "whisper-1"),
        "language": (None, "pt"),
    }
    headers = {"Authorization": f"Bearer {api_key}"}

    try:
        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(
                "https://api.openai.com/v1/audio/transcriptions",
                files=files,
                headers=headers,
            )
    except Exception as e:
        logger.error(f"Whisper request falhou: {e}")
        return f"(erro de rede na transcrição: {e})"

    if response.status_code == 401:
        return "(Whisper rejeitou a chave — OPENAI_API_KEY inválida)"
    if response.status_code == 429:
        return "(Whisper rate limited — tenta de novo em 30s)"
    if response.status_code != 200:
        logger.error(f"Whisper {response.status_code}: {response.text[:200]}")
        return f"(erro Whisper {response.status_code})"

    try:
        transcricao = (response.json().get("text") or "").strip()
    except Exception as e:
        return f"(resposta inválida do Whisper: {e})"

    if not transcricao:
        return "(áudio sem fala detectável)"

    if len(_audio_cache) >= _CACHE_MAX:
        _audio_cache.pop(next(iter(_audio_cache)))
    _audio_cache[file_hash] = transcricao
    return transcricao


async def processar_xlsx(file_url: str, caption: str = "") -> list[dict]:
    """Extrai dados de .xlsx (Excel) como texto tabular."""
    from openpyxl import load_workbook  # import lazy

    data = await _download(file_url)
    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)

    partes = []
    for nome_sheet in wb.sheetnames:
        ws = wb[nome_sheet]
        partes.append(f"## Planilha: {nome_sheet}")
        for linha in ws.iter_rows(values_only=True):
            celulas = [str(v) if v is not None else "" for v in linha]
            linha_str = "\t".join(celulas).strip()
            if linha_str:
                partes.append(linha_str)
        partes.append("")

    texto = "\n".join(partes) if partes else "(planilha vazia)"
    prompt = caption or "Analise esta planilha."
    return [{"type": "text", "text": f"{prompt}\n\n---\n\n{texto}"}]
